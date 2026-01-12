from docx import Document

doc = Document()
doc.add_heading('ACUERDO DE COLABORACIÓN PARA TRABAJO FIN DE GRADO (TFG TÉCNICO)', level=1)

sections = [
    """
DATOS DE LAS PARTES

Lugar: ____________________________
Fecha: ____ / ____ / ______

Integrante 1
Nombre completo: ______________________________________
DNI / Pasaporte: ______________________________________
Grado / Universidad: _________________________________

Integrante 2
Nombre completo: ______________________________________
DNI / Pasaporte: ______________________________________
Grado / Universidad: _________________________________
""",
    """
OBJETO DEL ACUERDO

El presente acuerdo regula la colaboración técnica y académica para la realización conjunta de un Trabajo Fin de Grado de carácter técnico (software, ingeniería, sistemas, IA u otros).
""",
    """
DETERMINACIÓN DEL TEMA DEL TFG

El tema será decidido de forma conjunta mediante reuniones técnicas de análisis, viabilidad y alcance.

Tema provisional:
______________________________________________________________
""",
    """
PLAN DE TRABAJO TÉCNICO

Fases:
- Análisis de requisitos y estado del arte
- Diseño de arquitectura técnica
- Desarrollo e implementación
- Pruebas y validación
- Documentación técnica
- Preparación de la defensa

Distribución inicial de tareas:
Investigación técnica: ______________________
Desarrollo backend / frontend / IA: ______________________
Documentación técnica: ______________________
Integración y pruebas: ______________________
""",
    """
REPORTES PERIÓDICOS

Frecuencia:
☐ Semanal   ☐ Quincenal   ☐ Mensual

Cada reporte incluirá:
- Avances técnicos
- Problemas y riesgos
- Próximas tareas
- Tiempo invertido
""",
    """
REUNIONES DE SEGUIMIENTO

Modalidad:
☐ Presencial   ☐ Online

Frecuencia aproximada: ______________________
""",
    """
COMPROMISOS

Los integrantes se comprometen a:
- Cumplir tareas técnicas asignadas
- Mantener repositorios actualizados (Git)
- Documentar correctamente el código
- Comunicar incidencias técnicas
""",
    """
INCUMPLIMIENTO

La falta reiterada de avances técnicos, reportes o comunicación será considerada incumplimiento grave y podrá comunicarse al tutor.
""",
    """
PROPIEDAD DEL TRABAJO

El trabajo, código fuente y documentación serán de propiedad conjunta.
""",
    """
DURACIÓN DEL ACUERDO

Vigente desde la firma hasta la entrega y defensa del TFG.
""",
    """
FIRMAS

Integrante 1
Nombre:
Firma:
Fecha:

Integrante 2
Nombre:
Firma:
Fecha:
"""
]

for s in sections:
    doc.add_paragraph(s)

file_path = "/"
doc.save(file_path)

file_path
